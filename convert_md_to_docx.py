#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter for RideBuddy Documentation
Converts markdown files to professional Word documents with proper formatting
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def setup_document_styles(doc):
    """Setup professional document styles"""
    
    # Define custom styles
    styles = doc.styles
    
    # Title style
    if 'CustomTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Segoe UI'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0x2F, 0x54, 0x96)  # Professional blue
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Heading styles
    heading_sizes = [18, 16, 14, 12]
    for i, size in enumerate(heading_sizes, 1):
        style_name = f'CustomHeading{i}'
        if style_name not in [s.name for s in styles]:
            heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Segoe UI'
            heading_font.size = Pt(size)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0x1F, 0x36, 0x56)  # Dark blue
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    # Code style
    if 'CustomCode' not in [s.name for s in styles]:
        code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(9)
        code_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
    
    # Quote style
    if 'CustomQuote' not in [s.name for s in styles]:
        quote_style = styles.add_style('CustomQuote', WD_STYLE_TYPE.PARAGRAPH)
        quote_font = quote_style.font
        quote_font.name = 'Segoe UI'
        quote_font.size = Pt(10)
        quote_font.italic = True
        quote_font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.right_indent = Inches(0.5)
        quote_style.paragraph_format.space_before = Pt(6)
        quote_style.paragraph_format.space_after = Pt(6)

def add_table_of_contents(doc, content):
    """Extract and add table of contents"""
    
    # Find all headings
    headings = re.findall(r'^(#{1,4})\s+(.+)$', content, re.MULTILINE)
    
    if headings:
        # Add TOC title
        toc_title = doc.add_paragraph('Table of Contents', style='CustomHeading1')
        
        # Add TOC entries
        for level, title in headings:
            # Clean title (remove emojis and special characters)
            clean_title = re.sub(r'[^\w\s-]', '', title).strip()
            if clean_title:
                level_num = len(level)
                indent = (level_num - 1) * 0.3
                
                toc_entry = doc.add_paragraph()
                toc_entry.paragraph_format.left_indent = Inches(indent)
                
                run = toc_entry.add_run(clean_title)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                
                if level_num == 1:
                    run.font.bold = True
        
        # Add page break after TOC
        doc.add_page_break()

def parse_and_convert_markdown(content, doc):
    """Parse markdown content and convert to Word document"""
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_language = ""
    code_content = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                # Start of code block
                in_code_block = True
                code_language = line[3:].strip()
                code_content = []
            else:
                # End of code block
                in_code_block = False
                
                # Add code block to document
                if code_content:
                    # Add language label if specified
                    if code_language:
                        label_p = doc.add_paragraph(f"Code ({code_language}):", style='Normal')
                        label_p.runs[0].font.bold = True
                        label_p.runs[0].font.size = Pt(9)
                    
                    # Add code content
                    code_text = '\n'.join(code_content)
                    code_p = doc.add_paragraph(code_text, style='CustomCode')
                    
                    # Add border to code block
                    try:
                        pPr = code_p._element.get_or_add_pPr()
                        pBdr = OxmlElement('w:pBdr')
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:sz'), '4')
                            border.set(qn('w:space'), '1')
                            border.set(qn('w:color'), 'CCCCCC')
                            pBdr.append(border)
                        pPr.append(pBdr)
                    except:
                        pass  # Fallback if border styling fails
            
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2)
                
                # Clean title (remove emojis)
                clean_title = re.sub(r'[^\w\s\-\(\)\[\]\.\,\:\;\/]', '', title).strip()
                
                if level == 1:
                    style = 'CustomTitle' if i < 5 else 'CustomHeading1'  # First heading as title
                else:
                    style = f'CustomHeading{min(level, 4)}'
                
                doc.add_paragraph(clean_title, style=style)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50, style='Normal')
        
        # Handle bullet points
        elif line.strip().startswith(('- ', '* ', '+ ')):
            text = line.strip()[2:]
            # Clean text (remove emojis)
            clean_text = re.sub(r'[^\w\s\-\(\)\[\]\.\,\:\;\/\%\<\>\=\+]', '', text).strip()
            p = doc.add_paragraph(clean_text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            # Clean text (remove emojis)
            clean_text = re.sub(r'[^\w\s\-\(\)\[\]\.\,\:\;\/\%\<\>\=\+]', '', text).strip()
            p = doc.add_paragraph(clean_text, style='List Number')
        
        # Handle tables (simple conversion)
        elif '|' in line and line.strip().startswith('|'):
            table_lines = []
            j = i
            
            # Collect table lines
            while j < len(lines) and '|' in lines[j]:
                if not lines[j].strip().replace('|', '').replace('-', '').replace(' ', ''):
                    # Skip separator lines
                    j += 1
                    continue
                table_lines.append(lines[j])
                j += 1
            
            if table_lines:
                # Create table
                rows = len(table_lines)
                cols = len(table_lines[0].split('|')) - 2  # Remove empty first and last
                
                if cols > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, line in enumerate(table_lines):
                        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
                        for col_idx, cell_text in enumerate(cells[:cols]):
                            if row_idx < len(table.rows) and col_idx < len(table.columns):
                                # Clean cell text
                                clean_text = re.sub(r'[^\w\s\-\(\)\[\]\.\,\:\;\/\%\<\>\=\+]', '', cell_text).strip()
                                table.cell(row_idx, col_idx).text = clean_text
                
                i = j - 1  # Skip processed table lines
        
        # Handle block quotes
        elif line.strip().startswith('>'):
            quote_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('>'):
                quote_lines.append(lines[j].strip()[1:].strip())
                j += 1
            
            if quote_lines:
                quote_text = ' '.join(quote_lines)
                # Clean quote text
                clean_text = re.sub(r'[^\w\s\-\(\)\[\]\.\,\:\;\/\%\<\>\=\+]', '', quote_text).strip()
                doc.add_paragraph(clean_text, style='CustomQuote')
                i = j - 1
        
        # Handle regular paragraphs
        elif line.strip():
            # Clean text (remove emojis and excessive special characters)
            clean_text = re.sub(r'[^\w\s\-\(\)\[\]\.\,\:\;\/\%\<\>\=\+\*\#]', '', line).strip()
            
            if clean_text:
                # Handle bold and italic formatting
                p = doc.add_paragraph()
                
                # Split by formatting patterns
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', clean_text)
                
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                        # Italic text
                        run = p.add_run(part[1:-1])
                        run.font.italic = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Inline code
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    else:
                        # Normal text
                        if part.strip():
                            p.add_run(part)
        
        i += 1

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """Convert a markdown file to DOCX format"""
    
    print(f"Converting {md_file_path} to {docx_file_path}...")
    
    try:
        # Read markdown content
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Setup professional styles
        setup_document_styles(doc)
        
        # Add table of contents (optional, based on headings)
        # add_table_of_contents(doc, content)
        
        # Convert markdown content
        parse_and_convert_markdown(content, doc)
        
        # Add footer with generation info
        section = doc.sections[0]
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = f"Generated from {Path(md_file_path).name} - RideBuddy Pro Documentation"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        doc.save(docx_file_path)
        print(f"✅ Successfully converted to {docx_file_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error converting {md_file_path}: {e}")
        return False

def main():
    """Main conversion function"""
    
    print("🔄 RideBuddy Documentation Converter - Markdown to DOCX")
    print("=" * 60)
    
    # Define file pairs
    files_to_convert = [
        {
            'md': 'RIDEBUDDY_SYSTEM_DOCUMENTATION.md',
            'docx': 'RideBuddy_System_Documentation.docx'
        },
        {
            'md': 'INSTALLATION_AND_SETUP_GUIDE.md', 
            'docx': 'RideBuddy_Installation_Setup_Guide.docx'
        }
    ]
    
    successful_conversions = 0
    
    for file_pair in files_to_convert:
        md_path = file_pair['md']
        docx_path = file_pair['docx']
        
        if os.path.exists(md_path):
            if convert_markdown_to_docx(md_path, docx_path):
                successful_conversions += 1
        else:
            print(f"❌ Source file not found: {md_path}")
    
    print("\n" + "=" * 60)
    print(f"📊 Conversion Summary:")
    print(f"✅ Successful conversions: {successful_conversions}")
    print(f"📁 Total files processed: {len(files_to_convert)}")
    
    if successful_conversions == len(files_to_convert):
        print("🎉 All files converted successfully!")
    else:
        print("⚠️  Some conversions failed. Please check the error messages above.")
    
    print("\n📋 Generated DOCX files:")
    for file_pair in files_to_convert:
        if os.path.exists(file_pair['docx']):
            file_size = os.path.getsize(file_pair['docx']) / 1024  # KB
            print(f"   • {file_pair['docx']} ({file_size:.1f} KB)")

if __name__ == "__main__":
    main()